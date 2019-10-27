# -*- coding: utf-8 -*-
"""
Created on Wed Aug  8 09:50:05 2018

Python操作word常见方法示例【win32com与docx模块】
https://www.jb51.net/article/139679.htm

"""

import docx
doc = docx.Document()
table = doc.add_table(rows=1, cols=3, style='Medium Shading 2 Accent 6') #创建带边框的表格
hdr_cells = table.rows[0].cells # 获取第0行所有所有单元格
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# 添加三行数据
data_lines = 3
for i in range(data_lines):
  cells = table.add_row().cells
  cells[0].text = 'Name%s' % i
  cells[1].text = 'Id%s' % i
  cells[2].text = 'Desc%s' % i
rows = 2
cols = 4
table = doc.add_table(rows=rows, cols=cols)
val = 1
for i in range(rows):
  cells = table.rows[i].cells
  for j in range(cols):
    cells[j].text = str(val * 10)
    val += 1
doc.save('tmp.docx')